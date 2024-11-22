# Librerías para manejo de datos y cálculos
import pandas as pd  # Biblioteca utilizada para el manejo y análisis de datos estructurados en forma de DataFrames.
import numpy as np  # Biblioteca para cálculos matemáticos y operaciones con arreglos/matrices de gran tamaño.

# Librerías para visualización y análisis estadístico
import matplotlib.pyplot as plt  # Biblioteca para crear gráficos y visualizaciones de datos.
from scipy.stats import skew, kurtosis, linregress  # Módulo de Scipy para análisis estadístico:
# - skew: Calcula el sesgo (asimetría) de una distribución.
# - kurtosis: Calcula la curtosis (forma de las colas) de una distribución.
# - linregress: Realiza una regresión lineal para un conjunto de datos.

# Librería para generar documentos de texto en formato .docx
from docx import Document  # Permite crear y manipular documentos de Microsoft Word (.docx).
from docx.shared import Inches  # Proporciona herramientas para configurar tamaños de elementos como imágenes en documentos.

# Cargar el archivo CSV base de datos
ruta_archivo = r'C:\Users\cesar\OneDrive\Escritorio\DayannaProyectoEstadistica\calidad_de_agua.csv'
df = pd.read_csv(ruta_archivo)  # Lee el archivo CSV

# Seleccionar dos columnas correlacionadas (por ejemplo, 'pH' y 'Oxígeno disuelto')
variable_x = pd.to_numeric(df['pH'], errors='coerce').dropna()
variable_y = pd.to_numeric(df['OXIGENO DISUELTO (mg O2/l)'], errors='coerce').dropna()

# Asegurar que ambas variables tengan el mismo número de datos
variable_x = variable_x[:len(variable_y)]
variable_y = variable_y[:len(variable_x)]

# Función para generar gráficos y estadísticas de frecuencia
def generar_frecuencias(variable, nombre_variable):
    # Definir número de intervalos
    num_clases = int(np.sqrt(len(variable)))
    bins = np.linspace(variable.min(), variable.max(), num_clases + 1)
    hist, _ = np.histogram(variable, bins=bins)

    # Crear tabla de frecuencias
    frecuencias = pd.DataFrame({
        'Intervalo': [f"[{bins[i]:.2f} - {bins[i + 1]:.2f})" for i in range(len(bins) - 1)],
        'Frecuencia': hist
    })

    # Filtrar frecuencias mayores a 0
    frecuencias = frecuencias[frecuencias['Frecuencia'] > 0]

    # Histograma
    plt.figure(figsize=(8, 6))
    plt.hist(variable, bins=bins, color='blue', alpha=0.7, rwidth=0.85)
    plt.title(f"Histograma de {nombre_variable}")
    plt.xlabel(nombre_variable)
    plt.ylabel("Frecuencia")
    plt.grid(alpha=0.5)
    plt.savefig(f"histograma_{nombre_variable.replace(' ', '_')}.png")
    plt.close()

    # Polígono de frecuencia
    plt.figure(figsize=(8, 6))
    puntos_x = [(bins[i] + bins[i + 1]) / 2 for i in range(len(bins) - 1)]
    plt.plot(puntos_x, hist, marker='o', color='red', label="Polígono de frecuencia")
    plt.title(f"Polígono de frecuencia de {nombre_variable}")
    plt.xlabel(nombre_variable)
    plt.ylabel("Frecuencia")
    plt.grid(alpha=0.5)
    plt.legend()
    plt.savefig(f"poligono_{nombre_variable.replace(' ', '_')}.png")
    plt.close()

    # Gráfico de pastel
    plt.figure(figsize=(8, 8))
    plt.pie(hist, labels=[f"[{bins[i]:.2f} - {bins[i + 1]:.2f})" for i in range(len(bins) - 1)],
            autopct='%1.1f%%', startangle=140)
    plt.title(f"Gráfico de pastel de {nombre_variable}")
    plt.savefig(f"torta_{nombre_variable.replace(' ', '_')}.png")
    plt.close()

    return frecuencias

# Generar frecuencias, histogramas, polígonos y gráficos de pastel
frecuencias_x = generar_frecuencias(variable_x, "pH")
frecuencias_y = generar_frecuencias(variable_y, "Oxígeno Disuelto")

# Medidas estadísticas para variable_x
media_x = variable_x.mean()
mediana_x = variable_x.median()
moda_x = variable_x.mode()[0]
varianza_x = variable_x.var()
desviacion_estandar_x = variable_x.std()
coef_variacion_x = desviacion_estandar_x / media_x
asimetria_x = skew(variable_x)
curtosis_x = kurtosis(variable_x)

# Medidas estadísticas para variable_y
media_y = variable_y.mean()
mediana_y = variable_y.median()
moda_y = variable_y.mode()[0]
varianza_y = variable_y.var()
desviacion_estandar_y = variable_y.std()
coef_variacion_y = desviacion_estandar_y / media_y
asimetria_y = skew(variable_y)
curtosis_y = kurtosis(variable_y)

# Regresión lineal
slope, intercept, r_value, p_value, std_err = linregress(variable_x, variable_y)

# Generar diagrama de dispersión y línea de regresión
plt.figure(figsize=(10, 6))
plt.scatter(variable_x, variable_y, color='blue', alpha=0.7, label='Datos')
plt.plot(variable_x, slope * variable_x + intercept, color='red', label=f'Regresión lineal (R²={r_value**2:.2f})')
plt.title('Diagrama de Dispersión y Línea de Regresión')
plt.xlabel('pH')
plt.ylabel('Oxígeno Disuelto')
plt.legend()
plt.grid(alpha=0.5)
plt.savefig('diagrama_dispersión_regresion.png')
plt.close()

# Crear un documento Word
doc = Document()
doc.add_heading('Análisis Completo de Variables: pH y Oxígeno Disuelto', 0)

# Agregar introducción
doc.add_heading('Introducción', level=1)
doc.add_paragraph(
    "El presente análisis tiene como objetivo evaluar las características estadísticas y relaciones entre las variables este estudio fue realizado en el río Cauca. "
    "El pH y el Oxígeno Disuelto (mg O2/l) son fundamentales para entender el estado ecológico del agua y su capacidad para sustentar vida acuática."
)

# Agregar objetivos
doc.add_heading('Objetivos', level=1)
doc.add_heading('General', level=2)
doc.add_paragraph("Realizar un análisis estadístico completo de las variables seleccionadas para identificar patrones, relaciones y su relevancia en la calidad del agua.")
doc.add_heading('Específicos', level=2)
doc.add_paragraph("Describir las distribuciones de las variables mediante histogramas, polígonos de frecuencia y gráficos de pastel.")
doc.add_paragraph("Calcular y analizar las principales medidas de tendencia central y dispersión para ambas variables.")
doc.add_paragraph("Determinar la relación entre las variables mediante una regresión lineal y evaluar su correlación.")

# Agregar resultados de frecuencias y gráficos
doc.add_heading('Frecuencias y Representaciones de pH', level=1)
table_x = doc.add_table(rows=1, cols=2)
table_x.style = 'Table Grid'
hdr_cells = table_x.rows[0].cells
hdr_cells[0].text = 'Intervalos'
hdr_cells[1].text = 'Frecuencia'
for _, row in frecuencias_x.iterrows():
    row_cells = table_x.add_row().cells
    row_cells[0].text = row['Intervalo']
    row_cells[1].text = str(row['Frecuencia'])
doc.add_picture('histograma_pH.png', width=Inches(4))
doc.add_picture('poligono_pH.png', width=Inches(4))
doc.add_picture('torta_pH.png', width=Inches(4))

doc.add_heading('Frecuencias y Representaciones de Oxígeno Disuelto', level=1)
table_y = doc.add_table(rows=1, cols=2)
table_y.style = 'Table Grid'
hdr_cells = table_y.rows[0].cells
hdr_cells[0].text = 'Intervalos'
hdr_cells[1].text = 'Frecuencia'
for _, row in frecuencias_y.iterrows():
    row_cells = table_y.add_row().cells
    row_cells[0].text = row['Intervalo']
    row_cells[1].text = str(row['Frecuencia'])
doc.add_picture('histograma_Oxígeno_Disuelto.png', width=Inches(4))
doc.add_picture('poligono_Oxígeno_Disuelto.png', width=Inches(4))
doc.add_picture('torta_Oxígeno_Disuelto.png', width=Inches(4))

# Agregar medidas estadísticas
doc.add_heading('Medidas Estadísticas', level=1)
doc.add_heading('pH', level=2)
doc.add_paragraph(f"Media: {media_x:.2f}")
doc.add_paragraph(f"Mediana: {mediana_x:.2f}")
doc.add_paragraph(f"Moda: {moda_x:.2f}")
doc.add_paragraph(f"Varianza: {varianza_x:.2f}")
doc.add_paragraph(f"Desviación Estándar: {desviacion_estandar_x:.2f}")
doc.add_paragraph(f"Coeficiente de Variación: {coef_variacion_x:.2f}")
doc.add_paragraph(f"Asimetría: {asimetria_x:.2f}")
doc.add_paragraph(f"Curtosis: {curtosis_x:.2f}")

doc.add_heading('Oxígeno Disuelto', level=2)
doc.add_paragraph(f"Media: {media_y:.2f}")
doc.add_paragraph(f"Mediana: {mediana_y:.2f}")
doc.add_paragraph(f"Moda: {moda_y:.2f}")
doc.add_paragraph(f"Varianza: {varianza_y:.2f}")
doc.add_paragraph(f"Desviación Estándar: {desviacion_estandar_y:.2f}")
doc.add_paragraph(f"Coeficiente de Variación: {coef_variacion_y:.2f}")
doc.add_paragraph(f"Asimetría: {asimetria_y:.2f}")
doc.add_paragraph(f"Curtosis: {curtosis_y:.2f}")

# Agregar fórmula de la regresión lineal
doc.add_heading('Ecuación de la Regresión Lineal', level=1)
doc.add_paragraph(f"Ecuación de la Regresión Lineal: y = {slope:.2f}x + {intercept:.2f}")

# Guardar el documento
doc.save("análisis_completo.docx")

# Imprimir mensaje de éxito
print("El análisis ha sido completado y guardado en 'análisis_completo.docx'.")
